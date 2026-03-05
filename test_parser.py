"""
Парсер Word-тестов → DataFrame для doc_testanswers
Запускать в Spyder: выделить нужный блок → F9 или Run Selection

Поддерживаемые форматы:
  TABLE_MODE = True  — Таблица (одна колонка):
      строка 1 — «1. Текст вопроса?»     (жирный)
      строки 2+ — «А) текст варианта»    (правильный — красный)
      пустая строка между вопросами — игнорируется

  TABLE_MODE = False — Параграфы (старый формат):
      жирный параграф = вопрос
      строки А)/Б)/В)/Г) = варианты

Архитектура БД (актуальная):
  options         — JSON: [{"id": "a", "text": "..."}, ...]
  correct_answers — JSON: ["a"]   (может быть несколько)
  a, b, c, d     — text (обратная совместимость, до 4 вариантов)
  answer          — text (первый правильный, обратная совместимость)
"""

import re
import json
import pandas as pd
from docx import Document
from docx.oxml.ns import qn

import os
import sqlalchemy as sqla

pd.options.mode.chained_assignment = None

# ── Подключение к БД ──────────────────────────────────────────────────────────
servername = 'postgresql://vinkom:a23sSD$sdf@192.168.70.221:5435/vinkom_data1'
engine = sqla.create_engine(servername, connect_args={'connect_timeout': 30})


# ── Настройки ─────────────────────────────────────────────────────────────────

DOCX_FILE     = r'добыча урана.docx'   # ← путь к файлу
CATEGORY_CODE = 'uran'                 # ← код категории (ilc_categories)
OUTPUT_FILE   = r'result_uran.xlsx'   # ← куда сохранить ('' = не сохранять)
TABLE_MODE    = True                   # True = таблица, False = параграфы


# ── Регулярки ─────────────────────────────────────────────────────────────────

# Вариант ответа: начинается с кириллической или латинской буквы + разделитель
# \s включает обычный пробел, но НЕ \xa0 — добавляем его явно в класс
OPTION_LINE_RE = re.compile(
    r'^[\s\xa0]*([АБВГДЕЖЗИABCDEFGHабвгдежзиabcdefgh])[\s\xa0]*[).\-–:][\s\xa0]*(.+)',
    re.UNICODE | re.DOTALL
)

# Нумерация вопроса в начале строки: «1.», «2)», «Вопрос 3.» и т.п.
QUESTION_NUM_RE = re.compile(r'^[\s\xa0]*(?:Вопрос[\s\xa0]*)?\d+[\s\xa0]*[.)]', re.IGNORECASE)


def normalise(text):
    """Заменяет неразрывные и нулевые пробелы на обычные, убирает лишние пробелы."""
    return text.replace('\xa0', ' ').replace('\u200b', '').replace('\u00ad', '').strip()

# Слипшийся «вопрос   А) вариант» в одной строке
INLINE_OPTION_RE = re.compile(
    r'^(.+?)\s{2,}([АБВГДЕЖЗИABCDEFGHабвгдежзиabcdefgh])\s*[).\-–:]\s*(.+)$',
    re.UNICODE | re.DOTALL
)


# ── Вспомогательные функции ───────────────────────────────────────────────────

def idx_to_letter(i):
    """0→'a', 1→'b', ..."""
    return chr(ord('a') + i)


def is_red(color_hex):
    if not color_hex:
        return False
    h = color_hex.strip('#').upper()
    if len(h) == 6:
        try:
            r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
            return r > 150 and g < 120 and b < 120
        except ValueError:
            pass
    return False


def get_para_info(para):
    """Возвращает (full_text, has_red, run_map)."""
    texts, run_map = [], []
    pos = 0
    for run in para.runs:
        if not run.text:
            continue
        rpr = run._r.find(qn('w:rPr'))
        red = False
        if rpr is not None:
            color_el = rpr.find(qn('w:color'))
            if color_el is not None and is_red(color_el.get(qn('w:val'))):
                red = True
        run_map.append((pos, pos + len(run.text), red))
        texts.append(run.text)
        pos += len(run.text)
    full_text = ''.join(texts)
    has_red = any(r[2] for r in run_map)
    return full_text, has_red, run_map


def segment_is_red(run_map, seg_start, seg_end):
    for rs, re_, rr in run_map:
        if rr and rs < seg_end and re_ > seg_start:
            return True
    return False


def cell_text_and_red(cell):
    """
    Возвращает (полный текст ячейки, is_red).
    Ячейка может содержать несколько параграфов — объединяем.
    Красный = хотя бы один ран во всех параграфах красный.
    """
    parts = []
    any_red = False
    for para in cell.paragraphs:
        full_text, has_red, _ = get_para_info(para)
        if full_text.strip():
            parts.append(full_text)
        if has_red:
            any_red = True
    return ' '.join(parts).strip(), any_red


# ── Общая функция сборки вопроса ──────────────────────────────────────────────

def build_question(category_code, question_text, cur_options):
    """
    cur_options: [(orig_letter, text, is_red), ...]
    Возвращает dict или None если вопрос невалиден.
    """
    if not question_text or len(cur_options) < 1:
        if question_text:
            print(f"  ⚠️  Нет вариантов: «{question_text[:65]}»")
        return None

    correct_indices = [idx for idx, (_, _, is_r) in enumerate(cur_options) if is_r]

    if not correct_indices:
        print(f"  ⚠️  Нет красного ответа: «{question_text[:65]}»")
        return None

    options_list = [
        {"id": idx_to_letter(idx), "text": opt_text}
        for idx, (_, opt_text, _) in enumerate(cur_options)
    ]
    correct_answers = [idx_to_letter(idx) for idx in correct_indices]

    if len(correct_indices) > 1:
        print(f"  ℹ️  Множественный ответ ({len(correct_indices)} правильных): «{question_text[:55]}»")

    abcd = {chr(ord('a') + i): cur_options[i][1] if i < len(cur_options) else ''
            for i in range(4)}

    return {
        'ilc_categories': category_code,
        'question':        question_text,
        'options':         json.dumps(options_list, ensure_ascii=False),
        'correct_answers': json.dumps(correct_answers, ensure_ascii=False),
        'a': abcd['a'], 'b': abcd['b'], 'c': abcd['c'], 'd': abcd['d'],
        'answer': correct_answers[0],
    }


# ── Парсер: ТАБЛИЧНЫЙ формат ──────────────────────────────────────────────────

def parse_docx_table(filepath, category_code):
    """
    Формат: таблица(ы) с одной колонкой (или первой значимой колонкой).
    Строка = вопрос если не является вариантом ответа.
    Строка = вариант если текст начинается с А)/Б)/В)/Г) и т.п.
    Красный цвет → правильный ответ.
    """
    doc = Document(filepath)
    questions   = []
    cur_q       = None
    cur_options = []

    def flush():
        nonlocal cur_q, cur_options
        rec = build_question(category_code, cur_q, cur_options)
        if rec:
            questions.append(rec)
        cur_q = None
        cur_options = []

    seen_cells = set()  # дедупликация объединённых ячеек (merged cells)

    for table in doc.tables:
        for row in table.rows:
            # Берём первую УНИКАЛЬНУЮ непустую ячейку строки
            # (объединённые ячейки python-docx возвращает несколько раз)
            cell_text, cell_red = '', False
            for cell in row.cells:
                cell_id = id(cell._tc)  # уникальный XML-элемент ячейки
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                t, r = cell_text_and_red(cell)
                if t:
                    cell_text, cell_red = t, r
                    break

            # Нормализуем: убираем \xa0 и невидимые символы перед матчингом
            cell_text = normalise(cell_text)
            if not cell_text:
                continue  # пустая строка — разделитель, пропускаем

            m = OPTION_LINE_RE.match(cell_text)
            if m:
                # Это вариант ответа (начинается с А)/Б)/В)/Г) и т.п.)
                opt_text = normalise(m.group(2))
                cur_options.append((m.group(1), opt_text, cell_red))

            elif QUESTION_NUM_RE.match(cell_text):
                # Явно пронумерованный вопрос → сбрасываем предыдущий
                flush()
                cur_q = normalise(QUESTION_NUM_RE.sub('', cell_text))

            else:
                # Нет буквенного префикса и нет номера → продолжение
                if cur_options:
                    # Дописываем к тексту последнего варианта
                    letter, prev_text, prev_red = cur_options[-1]
                    cur_options[-1] = (letter,
                                       prev_text + ' ' + cell_text,
                                       prev_red or cell_red)
                elif cur_q is not None:
                    # Дописываем к тексту вопроса
                    cur_q = cur_q + ' ' + cell_text
                else:
                    # Самое начало — вопрос без номера
                    cur_q = cell_text

    flush()  # последний вопрос

    df = pd.DataFrame(questions, columns=[
        'ilc_categories', 'question',
        'options', 'correct_answers',
        'a', 'b', 'c', 'd', 'answer'
    ])
    return df


# ── Парсер: ПАРАГРАФНЫЙ формат (старый) ──────────────────────────────────────

def para_to_lines(para):
    full_text, has_red, run_map = get_para_info(para)
    stripped = full_text.strip()
    if not stripped:
        return []

    results = []
    for line in stripped.split('\n'):
        line = line.strip()
        if not line:
            continue
        if OPTION_LINE_RE.match(line):
            try:
                pos = full_text.index(line[0])
                line_red = segment_is_red(run_map, pos, pos + len(line))
            except ValueError:
                line_red = has_red
            results.append((line, line_red))
            continue
        m = INLINE_OPTION_RE.match(line)
        if m:
            question_part = m.group(1).strip()
            option_letter = m.group(2)
            option_text   = m.group(3).strip()
            try:
                split_pos = full_text.index(option_letter, len(m.group(1)))
                q_red  = segment_is_red(run_map, 0, split_pos)
                op_red = segment_is_red(run_map, split_pos, len(full_text))
            except ValueError:
                q_red = op_red = has_red
            if question_part:
                results.append((question_part, q_red))
            results.append((f"{option_letter}) {option_text}", op_red))
            continue
        results.append((line, has_red))
    return results


def collect_lines(doc):
    raw = []
    for para in doc.paragraphs:
        raw.extend(para_to_lines(para))
    merged = []
    i = 0
    while i < len(raw):
        text, red = raw[i]
        while (i + 1 < len(raw)
               and not OPTION_LINE_RE.match(raw[i][0])
               and not OPTION_LINE_RE.match(raw[i + 1][0])
               and raw[i + 1][0].strip()):
            i += 1
            next_text, next_red = raw[i]
            if QUESTION_NUM_RE.match(next_text):
                merged.append((text, red))
                text, red = next_text, next_red
                break
            text = text.rstrip() + ' ' + next_text.lstrip()
            red  = red or next_red
        merged.append((text, red))
        i += 1
    return merged


def parse_docx_paragraphs(filepath, category_code):
    doc   = Document(filepath)
    lines = collect_lines(doc)
    questions   = []
    cur_q       = None
    cur_options = []

    def flush():
        nonlocal cur_q, cur_options
        rec = build_question(category_code, cur_q, cur_options)
        if rec:
            questions.append(rec)
        cur_q = None
        cur_options = []

    for text, has_red in lines:
        text = text.strip()
        if not text:
            continue
        m = OPTION_LINE_RE.match(text)
        if m:
            cur_options.append((m.group(1), m.group(2).strip(), has_red))
            continue
        flush()
        cur_q = QUESTION_NUM_RE.sub('', text).strip()

    flush()

    df = pd.DataFrame(questions, columns=[
        'ilc_categories', 'question',
        'options', 'correct_answers',
        'a', 'b', 'c', 'd', 'answer'
    ])
    return df


# ── Запуск ────────────────────────────────────────────────────────────────────

if TABLE_MODE:
    df = parse_docx_table(DOCX_FILE, CATEGORY_CODE)
else:
    df = parse_docx_paragraphs(DOCX_FILE, CATEGORY_CODE)

df

# Статистика
print(f"\n✅ Распарсено вопросов: {len(df)}")
if len(df):
    print(f"\nРаспределение ответов:\n{df['answer'].value_counts().sort_index()}")
    print(f"\nПервые 3 строки:")
    print(df[['ilc_categories', 'question', 'correct_answers', 'answer']].head(3).to_string())

# Сохранение
if OUTPUT_FILE and len(df):
    if OUTPUT_FILE.endswith('.xlsx'):
        df.to_excel(OUTPUT_FILE, index=False)
    else:
        df.to_csv(OUTPUT_FILE, index=False, encoding='utf-8-sig')
    print(f"\n💾 Сохранено: {OUTPUT_FILE}")


# ── Запись в БД ───────────────────────────────────────────────────────────────
df.to_sql('doc_testanswers', engine, if_exists='append', index=False)
print(f"\n✅ Записано в doc_testanswers: {len(df)} строк")


# ── Отладка: найти конкретный вопрос в файле ─────────────────────────────────
# doc = Document(DOCX_FILE)
# keyword = 'принципам гражданской'
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             t, r = cell_text_and_red(cell)
#             if keyword in t:
#                 print(f"[red={r}] {repr(t[:150])}")