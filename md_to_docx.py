import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            
        # Horizontal Rule
        elif line == '---':
            doc.add_page_break() # Or just a separator, but page break is cleaner for sections
            
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            parse_inline(p, line[2:])
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
            parse_inline(p, text)
            
        # Paragraphs
        elif line:
            p = doc.add_paragraph()
            parse_inline(p, line)
        else:
            # Empty line
            pass

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

def parse_inline(paragraph, text):
    # Very basic bold parser **text**
    # We clear paragraph if we are adding runs
    # But here we just append. For a more robust one we would split by regex
    
    # Clean up markers if we are just doing simple append
    # Let's do a slightly better regex split
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    md_to_docx('interfaces_report.md', 'interfaces_report.docx')
